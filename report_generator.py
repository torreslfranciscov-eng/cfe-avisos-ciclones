"""
report_generator.py
Generador automático de reportes Word (.docx / .docm) de Avisos de Ciclones Tropicales CFE.
Soporta Océano Pacífico y Océano Atlántico.
"""

import os
import re
import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

MESES_ES = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
    5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
    9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
}

FONT_NAME = "Noto Sans"
COLOR_PRIMARY = RGBColor(0x1E, 0x5B, 0x4F)   # Verde institucional CFE / SMN
COLOR_TITULAR = RGBColor(0x40, 0x31, 0x52)   # Morado / Vino
COLOR_TEXT = RGBColor(0x21, 0x25, 0x29)      # Gris oscuro / Negro


def set_run_style(run, text, font_name=FONT_NAME, size_pt=9, bold=False, italic=False, color=None):
    """Aplica formato tipográfico a un run de texto."""
    run.text = text
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def set_paragraph_text(p, text, font_name=FONT_NAME, size_pt=9, bold=False, italic=False, color=None, align=None):
    """Reemplaza el contenido de un párrafo manteniendo o aplicando el estilo deseado."""
    p.text = ""
    run = p.add_run()
    set_run_style(run, text, font_name=font_name, size_pt=size_pt, bold=bold, italic=italic, color=color)
    if align is not None:
        p.alignment = align
    return p


def format_fecha_hora(hora_gmt_str):
    """
    Formatea la fecha y hora a partir del texto de SMN.
    Retorna (fecha_texto, hora_texto, fecha_corta, hora_corta).
    """
    now = datetime.datetime.now()
    dia = now.day
    mes = MESES_ES.get(now.month, "septiembre")
    anio = now.year
    
    hora_local_match = re.search(r'(\d{1,2}:\d{2})', hora_gmt_str)
    if hora_local_match:
        hora_str = hora_local_match.group(1)
    else:
        hora_str = f"{now.hour:02d}:{now.minute:02d}"

    fecha_larga = f"{dia:02d} de {mes} de {anio} {hora_str} hrs"
    fecha_corta = f"{dia:02d} {mes} {anio}"
    hora_corta = f"{hora_str.replace(':', '.')}hrs"
    
    return fecha_larga, hora_str, fecha_corta, hora_corta


def generate_word_report(cyclone_data, template_path="plantilla_aviso.docx", output_dir="reportes_generados"):
    """
    Genera el reporte Word oficial a partir de la plantilla y los datos del ciclón.
    """
    os.makedirs(output_dir, exist_ok=True)
    doc = Document(template_path)

    # 1. Párrafo 0: Título Cuenca (Pacífico o Atlántico)
    p0 = doc.paragraphs[0]
    cuenca_nombre = cyclone_data.get("cuenca", "Océano Pacífico")
    set_paragraph_text(p0, f"Aviso de Ciclón Tropical en el {cuenca_nombre}", size_pt=9, bold=True, color=COLOR_PRIMARY)

    # 2. Párrafo 1: Fecha y Hora
    p1 = doc.paragraphs[1]
    hora_local_raw = cyclone_data["condiciones"].get("hora_local_gmt", "")
    fecha_larga, hora_str, fecha_corta, hora_corta = format_fecha_hora(hora_local_raw)
    set_paragraph_text(p1, fecha_larga, size_pt=9, bold=True, color=COLOR_PRIMARY)

    # 3. Párrafo 2: Titular del Aviso
    p2 = doc.paragraphs[2]
    titular = cyclone_data.get("titular", "").upper()
    if not titular:
        sistema = cyclone_data.get("sistema", "").upper()
        titular = f"{sistema} EN VIGILANCIA EN EL {cuenca_nombre.upper()}"
    set_paragraph_text(p2, titular, size_pt=9, bold=True, color=COLOR_TITULAR)

    # 4. Párrafo 4 & 5: Situación Actual
    p4 = doc.paragraphs[4]
    set_paragraph_text(p4, "SITUACIÓN ACTUAL", size_pt=9, bold=True, color=COLOR_TEXT)

    p5 = doc.paragraphs[5]
    situacion = cyclone_data.get("situacion_actual", "")
    if not situacion:
        situacion = f"{cyclone_data.get('sistema', 'El sistema')} se localiza en el {cuenca_nombre}."
    set_paragraph_text(p5, situacion, size_pt=9, bold=False, color=COLOR_TEXT)

    # 5. Tabla 0: Imágenes Satélite y Trayectoria
    tbl_img = doc.tables[0]
    cell_sat = tbl_img.cell(0, 0)
    cell_sat.text = ""
    p_sat = cell_sat.paragraphs[0]
    p_sat.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cyclone_data.get("img_sat_path") and os.path.exists(cyclone_data["img_sat_path"]):
        p_sat.add_run().add_picture(cyclone_data["img_sat_path"], width=Inches(3.3))

    cell_tray = tbl_img.cell(0, 1)
    cell_tray.text = ""
    p_tray = cell_tray.paragraphs[0]
    p_tray.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cyclone_data.get("img_tray_path") and os.path.exists(cyclone_data["img_tray_path"]):
        p_tray.add_run().add_picture(cyclone_data["img_tray_path"], width=Inches(3.3))

    # 6. Tabla 1: Condiciones Actuales
    tbl_cond = doc.tables[1]
    c = cyclone_data["condiciones"]

    hora_full = c.get("hora_local_gmt", "")
    if not "del" in hora_full and hora_full:
        hora_full = f"{hora_full} del {fecha_corta}"
    set_paragraph_text(tbl_cond.cell(1, 1).paragraphs[0], hora_full, size_pt=9)

    lat = c.get("latitud_norte", "--")
    lon = c.get("longitud_oeste", "--")
    set_paragraph_text(tbl_cond.cell(2, 1).paragraphs[0], f"Latitud Norte: {lat}", size_pt=9)
    set_paragraph_text(tbl_cond.cell(2, 2).paragraphs[0], f"Longitud Oeste: {lon}", size_pt=9)

    dist = c.get("distancia_costa", "--")
    set_paragraph_text(tbl_cond.cell(3, 1).paragraphs[0], dist, size_pt=9)

    desp = c.get("desplazamiento", "--")
    set_paragraph_text(tbl_cond.cell(4, 1).paragraphs[0], desp, size_pt=9)

    v_sost = c.get("vientos_sostenidos", "--")
    v_rach = c.get("vientos_rachas", "--")
    set_paragraph_text(tbl_cond.cell(5, 1).paragraphs[0], f"Sostenidos: {v_sost}", size_pt=9)
    set_paragraph_text(tbl_cond.cell(5, 2).paragraphs[0], f"Rachas: {v_rach}", size_pt=9)

    pres = c.get("presion_minima", "--")
    set_paragraph_text(tbl_cond.cell(6, 1).paragraphs[0], str(pres), size_pt=9)

    lluvia = c.get("pronostico_lluvia", "No se emite pronóstico de lluvia.")
    if not lluvia: lluvia = "No se emite pronóstico de lluvia."
    set_paragraph_text(tbl_cond.cell(7, 1).paragraphs[0], lluvia, size_pt=9)

    coment = c.get("comentarios_adicionales", "No representa peligro para el territorio mexicano.")
    if not coment: coment = "No representa peligro para el territorio mexicano."
    set_paragraph_text(tbl_cond.cell(8, 1).paragraphs[0], coment, size_pt=9)

    recom = c.get("recomendaciones", "No se emiten recomendaciones debido a su distancia.")
    if not recom: recom = "No se emiten recomendaciones debido a su distancia."
    set_paragraph_text(tbl_cond.cell(9, 1).paragraphs[0], recom, size_pt=9)

    # 7. Párrafo 6: Próximo Aviso
    p6 = doc.paragraphs[6]
    prox_aviso = cyclone_data.get("proximo_aviso", "").upper()
    if not prox_aviso:
        prox_aviso = "EL SIGUIENTE AVISO SE EMITIRÁ A LAS 03:15 HORAS TIEMPO DEL CENTRO O ANTES SI OCURREN CAMBIOS SIGNIFICATIVOS"
    set_paragraph_text(p6, prox_aviso, size_pt=8, bold=True, color=COLOR_TEXT)

    # 8. Nombre de archivo estándar CFE
    # e.g.: Aviso de Ciclón Océano Atlántico_D.T. EDOUARD_01 septiembre 2026_21.00hrs.docx
    sistema_raw = cyclone_data.get("nombre_limpio", cyclone_data.get("sistema", "CICLON")).strip()
    prefix = ""
    if "hurac" in sistema_raw.lower():
        nombre = re.sub(r'hurac[aá]n\s*', '', sistema_raw, flags=re.I).strip().upper()
        prefix = f"H. {nombre}"
    elif "tormenta" in sistema_raw.lower():
        nombre = re.sub(r'tormenta\s+tropical\s*', '', sistema_raw, flags=re.I).strip().upper()
        prefix = f"T.T. {nombre}"
    elif "depresi" in sistema_raw.lower():
        nombre = re.sub(r'depresi[oó]n\s+tropical\s*', '', sistema_raw, flags=re.I).strip().upper()
        prefix = f"D.T. {nombre}"
    else:
        prefix = sistema_raw.upper()

    filename = f"Aviso de Ciclón {cuenca_nombre}_{prefix}_{fecha_corta}_{hora_corta}.docx"
    filename = re.sub(r'[\\/*?:"<>|]', "", filename)
    output_docx_path = os.path.join(output_dir, filename)

    doc.save(output_docx_path)
    print(f"[OK] Reporte generado exitosamente: {output_docx_path}")
    return output_docx_path


if __name__ == "__main__":
    from smn_scraper import get_active_cyclones, fetch_cyclone_data

    print("=== Generando reportes de todos los ciclones activos (Pacífico y Atlántico) ===")
    cyclones = get_active_cyclones()
    for c in cyclones:
        data = fetch_cyclone_data(c["aviso_id"], basin_key=c["basin_key"])
        if data:
            generate_word_report(data)
